"""
分布式文档转换系统 — Ubuntu 调度节点
= 健康检查 + 文档转换 + 智能路由（阶段 4）+ 裁剪范围提取（阶段 5，v1）

路由策略：
    .docx/.xlsx/.pptx/.txt/.csv/.html/.xml/.zip/.epub → 本地 markitdown
    .png/.jpg/.jpeg/.bmp/.tif/.tiff/.webp               → 转发 Win11 PaddleOCR
    .pdf                                                 → pdfplumber 分类
        ├── 文本型 (可提取文字层) → 本地 markitdown
        └── 扫描型 (无文字层)    → 转发 Win11 PaddleOCR

版面重建（阶段 6，基于 Win11 OCR 返回的行级 bbox 做几何启发式处理）：
    Win11 /convert 现在会在响应里附带 "lines"（每行文字 + bbox 四角点 + 置信度，
    多页场景带 page 字段）。本节点收到后不再直接把 OCR 拼好的纯文本当 content，
    而是基于行高（字号代理）、行间距、左缩进这几个几何特征做段落/标题分组：
        - 行高明显大于全页中位行高 → 判定为标题，按高度比例分 1~3 级 (# / ## / ###)
        - 相邻行的垂直间距明显大于中位行高 → 判定为段落分界，插入空行
        - 段首左边界明显比正文基线靠右 → 判定为缩进段落，用 Markdown 引用块 "> " 包裹
        - 同一段内的多行会合并成一行连续文本（中英文交界处智能补/不补空格）
    这不是真正的版面分析（不处理多栏、表格），阈值可用环境变量调整（见下方配置）。
    只在图片 OCR 和扫描型 PDF（含裁剪）两条走 Win11 的路径上生效；本地
    markitdown 路径不受影响。重建失败或没有 lines 数据时，自动回退到 Win11
    原样拼接的纯文本，不影响可用性。响应新增 layout_reconstructed 字段标注是否
    实际做了重建；重建成功时原始纯文本保留在 content_raw 里备用。

裁剪范围提取（可选，/convert 表单新增 crop 字段）：
    crop = '{"unit":"ratio","x":0.1,"y":0.15,"width":0.4,"height":0.3}'
    - x/y/width/height 均为 0-1 的比例值，相对原图/PDF 页面宽高，与渲染 DPI 无关
    - 图片：Pillow 按比例换算像素坐标后裁剪，再转发裁剪后的图片给 Win11 OCR
    - 文本型 PDF：pdfplumber 按比例换算 PDF point 坐标，直接裁剪提取文字层（不走 OCR）
    - 扫描型 PDF：PyMuPDF 按比例换算 PDF point 坐标，逐页裁剪渲染成图片，逐页转发 Win11 OCR 后拼接
    - office 格式（docx/xlsx 等）不支持区域裁剪，crop 字段会被忽略（响应里标注 crop_ignored）
    - v1 限定：一个裁剪框应用到 PDF 所有页面，不支持逐页不同区域

输出格式（可选，/convert 表单新增 output_format 字段，默认 "md"）：
    output_format = "md" | "txt" | "docx"
    - md   ：原样返回转换出的 Markdown/纯文本内容（content 字段），行为与之前一致
    - txt  ：对内容做轻量 Markdown 语法剥离后返回（content 字段），不做完整解析
    - docx ：将内容转换生成真正的 .docx 文件，以 base64 编码放入 file_base64 字段返回
             （此时不再返回 content 字段）。生成引擎：
             · 优先用系统 pandoc（markdown -> docx），支持表格/脚注/更完整的列表与
               标题语义；可选 DOCX_REFERENCE_TEMPLATE 指定参考模板 docx，套用其中
               的标题/正文/字体等样式，得到统一的"标准样式"输出。
             · pandoc 未安装或转换失败时，自动降级到内置的轻量 python-docx 解析器
               （仅支持标题/列表/加粗/斜体/行内代码/代码块）。
             · 响应会附带 docx_engine 字段说明实际使用的引擎。
             · 源文件本身已是 .docx 时，禁止再选择 output_format=docx（会先经
               markitdown 拆成 Markdown 再重建，原始排版/图片/表格样式必然丢失，
               没有意义），返回 400。
    响应新增字段：output_format, output_filename, mime_type（docx 额外有 file_base64, docx_engine）

启动方式：
    NODE_NAME=ubuntu NODE_ROLE=light PORT=5000 \
    WIN11_OCR_URL=http://192.168.0.81:5000 \
    DOCX_REFERENCE_TEMPLATE=/opt/doc-converter/templates/standard.docx \
    python3 app.py

依赖：
    pip install flask flask-cors psutil requests pdfplumber pillow pymupdf python-docx "markitdown[all]"
    系统包（docx 输出走 pandoc 引擎所需，可选但强烈建议安装）：
    apt install pandoc

接口契约 v1.3：
    /health             → status, node, role, cpu, mem, gpu_*, ocr_mode
    /health/gpu         → GPU 诊断详情
    /supported_formats  → 节点能力声明（调度器动态发现）
    /convert            → 统一转换入口，自动路由，支持可选 crop 字段与 output_format 字段
                          OCR 路径新增 layout_reconstructed / content_raw 字段（见上方版面重建说明）
"""

import base64
import io
import json
import os
import re
import shutil
import subprocess
import time
import threading
from flask import Flask, request, jsonify
from flask_cors import CORS
import psutil
import requests
import pdfplumber
import fitz  # PyMuPDF，用于扫描型 PDF 的裁剪渲染
from PIL import Image
from markitdown import MarkItDown
from docx import Document
from docx.shared import Pt

app = Flask(__name__)

# ---- CORS ----
CORS(app, resources={
    r"/health":             {"origins": "*"},
    r"/health/gpu":         {"origins": "*"},
    r"/supported_formats":  {"origins": "*"},
    r"/convert":            {"origins": "*"},
    r"/":                   {"origins": "*"},
})

# ---- 节点身份 ----
NODE_NAME = os.environ.get("NODE_NAME", "ubuntu")
NODE_ROLE = os.environ.get("NODE_ROLE", "light")
START_TIME = time.time()

# ---- 任务计数器（线程安全）----
_lock = threading.Lock()
_active_tasks = 0


def task_started():
    with _lock:
        global _active_tasks
        _active_tasks += 1


def task_finished():
    with _lock:
        global _active_tasks
        _active_tasks = max(0, _active_tasks - 1)


# ---- 调度配置 ----
WIN11_OCR_URL = os.environ.get("WIN11_OCR_URL", "http://192.168.0.81:5000")
MAX_FILE_SIZE_MB = int(os.environ.get("MAX_FILE_SIZE_MB", 50))
OCR_TIMEOUT = int(os.environ.get("OCR_TIMEOUT", 120))  # OCR 转发超时（秒）
PDF_TEXT_THRESHOLD = int(os.environ.get("PDF_TEXT_THRESHOLD", 50))  # PDF 分类：最少可提取字符数
CROP_PDF_ZOOM = float(os.environ.get("CROP_PDF_ZOOM", 2.0))  # 扫描型 PDF 裁剪渲染倍率（越高越清晰，越慢）
CROP_EPS = 1e-4  # 裁剪比例边界判断的浮点误差容忍

# ---- 版面重建（基于 Win11 返回的行级 bbox 做几何启发式段落/标题分组）----
LAYOUT_RECONSTRUCTION_ENABLED = os.environ.get(
    "LAYOUT_RECONSTRUCTION_ENABLED", "true"
).strip().lower() not in ("false", "0", "no")  # 总开关，出问题时可整体回退到纯文本拼接
LAYOUT_HEADING_HEIGHT_RATIO = float(os.environ.get("LAYOUT_HEADING_HEIGHT_RATIO", 1.3))   # 行高/中位行高 ≥ 此值判定为标题
LAYOUT_PARAGRAPH_GAP_RATIO = float(os.environ.get("LAYOUT_PARAGRAPH_GAP_RATIO", 1.4))     # 行间距/中位行高 ≥ 此值判定为新段落
LAYOUT_INDENT_RATIO = float(os.environ.get("LAYOUT_INDENT_RATIO", 0.5))                   # 左缩进/中位行高 ≥ 此值判定为缩进段落（引用块）

# ---- docx 生成引擎（pandoc 优先，python-docx 作为降级兜底）----
# DOCX_REFERENCE_TEMPLATE：可选，指向一个 .docx 模板文件路径。pandoc 会读取该模板里
# 的样式（字体、字号、颜色、页边距、标题样式等）套用到生成的文档上，从而得到"标准/
# 给定模板样式"的输出，而不是 pandoc 默认的朴素样式。
# 制作模板的方法：用 Word/WPS 打开任意 docx，调整好"标题1/标题2/正文/列表/代码"等
# 样式后另存为一个干净的 docx（内容可留空或留一两行示例），把路径填到这里即可。
DOCX_REFERENCE_TEMPLATE = os.environ.get("DOCX_REFERENCE_TEMPLATE", "").strip()
PANDOC_TIMEOUT = int(os.environ.get("PANDOC_TIMEOUT", 30))
_PANDOC_PATH = shutil.which("pandoc")  # 启动时探测一次；未安装则自动降级到 python-docx

# ---- 路由表 ----
# 直接走本地 markitdown 的格式
LOCAL_ONLY_EXTENSIONS = {
    ".docx", ".xlsx", ".pptx", ".txt", ".csv",
    ".html", ".htm", ".xml", ".json", ".md",
    ".zip", ".epub", ".rtf", ".odt", ".ods", ".odp",
}

# 直接转发 Win11 OCR 的格式
OCR_ONLY_EXTENSIONS = {
    ".png", ".jpg", ".jpeg", ".bmp",
    ".tif", ".tiff", ".webp",
}

# PDF 需要分类判断
PDF_EXTENSION = ".pdf"

# ---- MarkItDown 实例 ----
markitdown = MarkItDown()


# ============================================================
#  路由：健康检查
# ============================================================

@app.route("/health", methods=["GET"])
def health():
    cpu_percent = psutil.cpu_percent(interval=0.2)
    mem = psutil.virtual_memory()

    with _lock:
        active = _active_tasks

    return jsonify({
        "status": "UP",
        "node": NODE_NAME,
        "role": NODE_ROLE,
        "cpu": round(cpu_percent / 100, 4),
        "mem": round(mem.percent / 100, 4),
        "mem_available_mb": round(mem.available / 1024 / 1024, 1),
        "active_tasks": active,
        "uptime_seconds": round(time.time() - START_TIME, 1),
        "gpu_available": False,
        "gpu_usable": False,
        "ocr_mode": "none",
    }), 200


@app.route("/health/gpu", methods=["GET"])
def health_gpu():
    return jsonify({
        "node": NODE_NAME,
        "gpu_available": False,
        "gpu_usable": False,
        "message": "Ubuntu light node has no GPU. Use Win11 heavy node for OCR tasks.",
    }), 200


# ============================================================
#  路由：能力声明（调度器动态发现）
# ============================================================

@app.route("/supported_formats", methods=["GET"])
def supported_formats():
    """声明本节点能力，调度器可通过此端点动态发现，无需硬编码格式列表。"""
    return jsonify({
        "node": NODE_NAME,
        "role": NODE_ROLE,
        "engine": "markitdown",
        "formats": sorted(LOCAL_ONLY_EXTENSIONS),
        "max_file_size_mb": MAX_FILE_SIZE_MB,
        "capabilities": {
            "ocr": False,
            "pdf_text_extraction": True,
            "office_formats": True,
            "image_formats": False,
            "crop": True,  # 支持 /convert 的可选 crop 字段（比例裁剪，v1：单框应用到全部页）
            "output_formats": sorted(SUPPORTED_OUTPUT_FORMATS),  # /convert 的可选 output_format 字段
            "layout_reconstruction": LAYOUT_RECONSTRUCTION_ENABLED,  # OCR 路径是否基于 bbox 做段落/标题重建
            "docx_engine": "pandoc" if _PANDOC_PATH else "python-docx-fallback",
            "docx_reference_template": bool(DOCX_REFERENCE_TEMPLATE and os.path.isfile(DOCX_REFERENCE_TEMPLATE)),
        },
        # 下游节点（如果本节点作为调度入口）
        "downstream": {
            "ocr_node": WIN11_OCR_URL,
            "ocr_formats": sorted(OCR_ONLY_EXTENSIONS | {PDF_EXTENSION}),
        },
    }), 200


# ============================================================
#  路由：首页（前端 UI）
# ============================================================

@app.route("/", methods=["GET"])
def index():
    """返回文档转换前端页面。API 信息见 /health 与 /supported_formats。"""
    html_path = os.path.join(os.path.dirname(__file__), "html", "document_converter.html")
    try:
        with open(html_path, "r", encoding="utf-8") as f:
            return f.read(), 200, {"Content-Type": "text/html; charset=utf-8"}
    except FileNotFoundError:
        return jsonify({
            "error": "Frontend not deployed",
            "hint": "Place document_converter.html in html/ directory",
        }), 404


# ============================================================
#  辅助：PDF 分类
# ============================================================

def _classify_pdf(file_bytes: bytes) -> str:
    """
    判断 PDF 是文本型还是扫描型。
    返回 "text" 或 "scanned"。
    策略：用 pdfplumber 提取前 3 页文字，字符数 > 阈值 → text。
    """
    try:
        total_chars = 0
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            pages_to_check = min(3, len(pdf.pages))
            for i in range(pages_to_check):
                page_text = pdf.pages[i].extract_text()
                if page_text:
                    # 去掉空白后计数
                    total_chars += len(re.sub(r"\s+", "", page_text))

        if total_chars >= PDF_TEXT_THRESHOLD:
            return "text"
        return "scanned"
    except Exception:
        # pdfplumber 打不开 → 可能是纯图片 PDF，归类为 scanned
        return "scanned"


# ============================================================
#  辅助：转发到 Win11 OCR
# ============================================================

def _forward_to_win11(file_bytes: bytes, filename: str, include_lines: bool = True) -> tuple:
    """转发文件到 Win11 OCR 节点。返回 (response_dict, http_status)。
    include_lines 显式传给 Win11 的 /convert，不依赖对方的默认值 —— 即使
    Win11 那边以后改了默认值，这里的行为也不会跟着变。"""
    try:
        resp = requests.post(
            f"{WIN11_OCR_URL}/convert",
            files={"file": (filename, io.BytesIO(file_bytes))},
            params={"include_lines": "true" if include_lines else "false"},
            timeout=OCR_TIMEOUT,
        )
        # 透传 Win11 的 JSON 响应
        try:
            return resp.json(), resp.status_code
        except ValueError:
            return {
                "success": False,
                "error": f"Win11 OCR returned non-JSON response (HTTP {resp.status_code})",
                "filename": filename,
            }, 502
    except requests.exceptions.ConnectionError:
        return {
            "success": False,
            "error": f"Win11 OCR node unreachable at {WIN11_OCR_URL}",
            "filename": filename,
            "failover": "OCR node down — retry later or use text-based conversion if applicable",
        }, 503
    except requests.exceptions.Timeout:
        return {
            "success": False,
            "error": f"Win11 OCR timed out after {OCR_TIMEOUT}s",
            "filename": filename,
        }, 504


# ============================================================
#  辅助：检查 Win11 健康状态
# ============================================================

def _check_win11_health() -> dict:
    """快速探活 Win11 节点，返回健康信息或 None。"""
    try:
        resp = requests.get(f"{WIN11_OCR_URL}/health", timeout=5)
        return resp.json() if resp.ok else None
    except Exception:
        return None


# ============================================================
#  辅助：裁剪范围解析与应用
# ============================================================

def _parse_crop(raw: str) -> dict:
    """解析并校验 crop 表单字段。合法返回 {x,y,width,height}（比例值），非法抛 ValueError。"""
    try:
        data = json.loads(raw)
    except (TypeError, ValueError):
        raise ValueError("crop is not valid JSON")

    if not isinstance(data, dict):
        raise ValueError("crop must be a JSON object")
    if data.get("unit") != "ratio":
        raise ValueError('crop.unit must be "ratio"')

    out = {}
    for key in ("x", "y", "width", "height"):
        val = data.get(key)
        if not isinstance(val, (int, float)) or isinstance(val, bool):
            raise ValueError(f"crop.{key} must be a number")
        if val < 0 or val > 1:
            raise ValueError(f"crop.{key} must be between 0 and 1")
        out[key] = float(val)

    if out["width"] <= 0 or out["height"] <= 0:
        raise ValueError("crop.width and crop.height must be > 0")
    if out["x"] + out["width"] > 1 + CROP_EPS:
        raise ValueError("crop.x + crop.width exceeds 1.0")
    if out["y"] + out["height"] > 1 + CROP_EPS:
        raise ValueError("crop.y + crop.height exceeds 1.0")

    return out


def _crop_image_bytes(file_bytes: bytes, crop: dict) -> bytes:
    """按比例裁剪图片，返回裁剪后的图片字节（保留原格式）。"""
    img = Image.open(io.BytesIO(file_bytes))
    w, h = img.size
    box = (
        round(crop["x"] * w),
        round(crop["y"] * h),
        round((crop["x"] + crop["width"]) * w),
        round((crop["y"] + crop["height"]) * h),
    )
    # 保证裁剪框至少 1x1 像素，避免比例太小时 Pillow 报错
    box = (
        min(box[0], w - 1),
        min(box[1], h - 1),
        max(box[2], box[0] + 1),
        max(box[3], box[1] + 1),
    )
    cropped = img.crop(box)
    buf = io.BytesIO()
    fmt = (img.format or "PNG")
    cropped.save(buf, format=fmt)
    return buf.getvalue()


def _crop_pdf_text(file_bytes: bytes, crop: dict) -> str:
    """文本型 PDF：按比例裁剪每一页并直接提取文字层，不走 OCR。"""
    parts = []
    with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
        for page in pdf.pages:
            w, h = page.width, page.height
            bbox = (
                crop["x"] * w,
                crop["y"] * h,
                (crop["x"] + crop["width"]) * w,
                (crop["y"] + crop["height"]) * h,
            )
            try:
                text = page.crop(bbox).extract_text() or ""
            except Exception:
                text = ""
            if text.strip():
                parts.append(text.strip())
    return "\n\n".join(parts)


def _crop_pdf_page_images(file_bytes: bytes, crop: dict) -> list:
    """扫描型 PDF：按比例裁剪每一页并渲染成图片。返回 [(page_no, png_bytes), ...]。"""
    images = []
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    try:
        mat = fitz.Matrix(CROP_PDF_ZOOM, CROP_PDF_ZOOM)
        for i, page in enumerate(doc, start=1):
            rect = page.rect
            clip = fitz.Rect(
                crop["x"] * rect.width,
                crop["y"] * rect.height,
                (crop["x"] + crop["width"]) * rect.width,
                (crop["y"] + crop["height"]) * rect.height,
            )
            pix = page.get_pixmap(matrix=mat, clip=clip)
            images.append((i, pix.tobytes("png")))
    finally:
        doc.close()
    return images


def _extract_ocr_text(result: dict) -> str:
    """从 Win11 OCR 响应里尽量取出文字内容，兼容字段名不完全确定的情况。"""
    for key in ("content", "text", "markdown"):
        val = result.get(key)
        if isinstance(val, str) and val.strip():
            return val.strip()
    return ""


# ============================================================
#  版面重建：基于 Win11 返回的行级 bbox 做几何启发式段落/标题分组
# ============================================================
# 思路：OCR 只给出逐行文字 + 四角点坐标，丢失了段落结构。这里不做真正的版面
# 分析（不处理多栏、表格），只用三个廉价的几何信号：
#   1. 行高（bbox 高度，字号的代理指标）明显偏大 → 标题
#   2. 相邻两行的垂直间距明显大于正文行高 → 段落分界
#   3. 段落起始左边界明显比正文基线靠右 → 缩进段落（引用块）
# 所有阈值都以"全页中位行高"为单位换算，而不是绝对像素值，这样不同 DPI/
# 分辨率渲染出来的图片都能用同一套默认阈值。

_ASCII_WORD_RE = re.compile(r'[A-Za-z0-9]')


def _smart_join(prev_text: str, next_text: str) -> str:
    """把同一段落里换行断开的两行文字拼回一句。
    只有断行两侧都是 ASCII 字母/数字（英文单词场景）才补一个空格，避免把单词
    粘在一起；中文/日文/韩文以及标点结尾的断行直接拼接，不引入多余空格。"""
    if not prev_text:
        return next_text
    if not next_text:
        return prev_text
    if _ASCII_WORD_RE.match(prev_text[-1]) and _ASCII_WORD_RE.match(next_text[0]):
        return prev_text + " " + next_text
    return prev_text + next_text


def _line_metrics(line: dict) -> dict:
    """从一行 OCR 结果的 bbox 里提取几何特征（顶/底/行高/左边界），供分组逻辑使用。"""
    xs = [pt[0] for pt in line["bbox"]]
    ys = [pt[1] for pt in line["bbox"]]
    top, bottom = min(ys), max(ys)
    return {
        "text": (line.get("text") or "").strip(),
        "top": top,
        "bottom": bottom,
        "height": max(bottom - top, 1e-6),
        "left": min(xs),
    }


def _reconstruct_page_layout(page_lines: list) -> str:
    """对单页（或单张图片）的行列表做段落/标题分组，返回重建后的 Markdown 文本。
    假设输入行已经大致按阅读顺序排列（PaddleOCR 默认从上到下），本函数不重新
    排序 —— 多栏排版下顺序可能是错的，这是已知局限，v1 不处理。"""
    metrics = [_line_metrics(ln) for ln in page_lines if ln.get("text", "").strip()]
    if not metrics:
        return ""

    heights = sorted(m["height"] for m in metrics)
    median_height = heights[len(heights) // 2] or 1.0

    # 正文基线左边界：只取行高接近中位值（即"正文行"，排除标题/异常行）的最小 left
    body_lefts = [m["left"] for m in metrics
                  if abs(m["height"] - median_height) <= 0.2 * median_height]
    baseline_left = min(body_lefts) if body_lefts else min(m["left"] for m in metrics)

    # ---- 分组：标题独占一组；正文按行间距切分段落 ----
    groups = []
    current = []
    for m in metrics:
        is_heading_line = m["height"] >= median_height * LAYOUT_HEADING_HEIGHT_RATIO
        if is_heading_line:
            if current:
                groups.append(current)
                current = []
            groups.append([m])
            continue

        if not current:
            current.append(m)
            continue

        gap = m["top"] - current[-1]["bottom"]
        indent_delta = abs(m["left"] - current[0]["left"])
        if (gap > median_height * LAYOUT_PARAGRAPH_GAP_RATIO
                or indent_delta > median_height * LAYOUT_INDENT_RATIO):
            groups.append(current)
            current = [m]
        else:
            current.append(m)

    if current:
        groups.append(current)

    # ---- 渲染：标题 → # 语法；缩进段落 → 引用块；普通段落 → 多行拼接成一段 ----
    rendered = []
    for group in groups:
        if len(group) == 1 and group[0]["height"] >= median_height * LAYOUT_HEADING_HEIGHT_RATIO:
            ratio = group[0]["height"] / median_height
            level = 1 if ratio >= 2.0 else (2 if ratio >= 1.6 else 3)
            rendered.append(f"{'#' * level} {group[0]['text']}")
            continue

        text = group[0]["text"]
        for m in group[1:]:
            text = _smart_join(text, m["text"])

        indent = group[0]["left"] - baseline_left
        if indent > median_height * LAYOUT_INDENT_RATIO:
            text = f"> {text}"

        rendered.append(text)

    return "\n\n".join(rendered)


def _reconstruct_layout(lines: list, force_page_markers: bool = False) -> str:
    """入口：把 Win11 返回的（可能跨页的）行列表重建成 Markdown 文本。
    按 line["page"]（缺省视为第 1 页）分组，页内重建，多页之间用
    "<!-- page N -->" 分隔（与本节点原有的裁剪多页拼接约定保持一致）。"""
    if not lines:
        return ""

    pages = {}
    for ln in lines:
        page_no = ln.get("page", 1)
        pages.setdefault(page_no, []).append(ln)

    use_markers = force_page_markers or len(pages) > 1
    parts = []
    for page_no in sorted(pages.keys()):
        page_md = _reconstruct_page_layout(pages[page_no])
        if use_markers:
            parts.append(f"<!-- page {page_no} -->" + (f"\n{page_md}" if page_md else ""))
        elif page_md:
            parts.append(page_md)

    return "\n\n".join(p for p in parts if p)


def _apply_layout_reconstruction(result: dict, force_page_markers: bool = False) -> None:
    """就地改写 OCR 响应 dict：能重建就把 content 换成重建后的 Markdown，
    原始纯文本挪到 content_raw；重建失败/没有行数据则原样不动。
    始终设置 layout_reconstructed 字段，方便下游判断是否发生了改写。"""
    result["layout_reconstructed"] = False

    if not LAYOUT_RECONSTRUCTION_ENABLED:
        result.pop("lines", None)
        return

    lines = result.get("lines")
    if not isinstance(lines, list) or not lines:
        result.pop("lines", None)
        return

    try:
        reconstructed = _reconstruct_layout(lines, force_page_markers=force_page_markers)
    except Exception as e:
        # 几何启发式本身出错不应该把整个转换搞失败 —— 记录原因，回退到原始文本
        result["layout_reconstruction_error"] = str(e)
        result.pop("lines", None)
        return

    if reconstructed.strip():
        result["content_raw"] = result.get("content", "")
        result["content"] = reconstructed
        result["layout_reconstructed"] = True

    result.pop("lines", None)


def _convert_ocr_cropped_pdf(file_bytes: bytes, filename: str, crop: dict,
                              t_start: float) -> dict:
    """扫描型 PDF + 裁剪：逐页裁剪渲染后转发 Win11 OCR，拼接结果。"""
    win11_health = _check_win11_health()
    if win11_health is None:
        return {
            "success": False,
            "error": "Win11 OCR node unreachable, cannot apply crop to scanned PDF",
            "filename": filename,
        }, 503

    try:
        page_images = _crop_pdf_page_images(file_bytes, crop)
    except Exception as e:
        return {
            "success": False,
            "error": f"Failed to render cropped PDF pages: {e}",
            "filename": filename,
        }, 500

    page_texts = []
    page_errors = []
    all_lines = []
    for page_no, img_bytes in page_images:
        result, status = _forward_to_win11(img_bytes, f"{filename}_page{page_no}.png")
        if result.get("success"):
            page_texts.append(f"<!-- page {page_no} -->\n{_extract_ocr_text(result)}")
            # 每张裁剪图对 Win11 来说都是独立的单页图片，它自己不会打 page 字段，
            # 这里用循环里真实的原始页码覆盖/补上，保证跨页重建时页码对得上。
            for ln in result.get("lines") or []:
                ln["page"] = page_no
                all_lines.append(ln)
        else:
            page_errors.append({"page": page_no, "error": result.get("error", f"HTTP {status}")})

    elapsed_ms = round((time.perf_counter() - t_start) * 1000)

    if not page_texts:
        return {
            "success": False,
            "error": "All cropped pages failed OCR",
            "filename": filename,
            "page_errors": page_errors,
            "route_elapsed_ms": elapsed_ms,
        }, 502

    resp = {
        "success": True,
        "filename": filename,
        "content": "\n\n".join(page_texts),
        "engine": "paddleocr_cpu",
        "routed_to": "win11_ocr",
        "crop_applied": True,
        "pages_processed": len(page_images),
        "route_elapsed_ms": elapsed_ms,
        "lines": all_lines,
    }
    # 裁剪场景本来就是逐页处理，页与页之间的分隔一直靠 "<!-- page N -->" 标记，
    # 这里强制沿用同样的约定，即使某次裁剪结果只剩一页也保留标记，行为不突变。
    _apply_layout_reconstruction(resp, force_page_markers=True)
    if page_errors:
        resp["page_errors"] = page_errors  # 部分页失败，但至少一页成功，仍算成功响应
    return resp, 200


# ============================================================
#  辅助：输出格式转换（md / txt / docx）
# ============================================================

SUPPORTED_OUTPUT_FORMATS = {"md", "txt", "docx"}

_INLINE_TOKEN_RE = re.compile(r'(\*\*\*.+?\*\*\*|\*\*.+?\*\*|\*.+?\*|__.+?__|_.+?_|`[^`]+?`)')


def _validate_output_format(raw: str) -> str:
    """校验 output_format 表单字段，非法抛 ValueError。"""
    fmt = (raw or "md").strip().lower()
    if fmt not in SUPPORTED_OUTPUT_FORMATS:
        raise ValueError(f'output_format must be one of {sorted(SUPPORTED_OUTPUT_FORMATS)}')
    return fmt


def _strip_markdown_light(md_text: str) -> str:
    """轻量剥离常见 Markdown 语法符号，得到更接近纯文本的内容（用于 txt 输出）。
    注意：这不是完整的 Markdown 解析器，只做常见符号的正则替换。"""
    text = md_text or ""
    text = re.sub(r'^(#{1,6})\s+', '', text, flags=re.MULTILINE)              # 标题符号
    text = re.sub(r'^```[^\n]*\n', '', text, flags=re.MULTILINE)              # 代码块起始围栏
    text = re.sub(r'^```\s*$', '', text, flags=re.MULTILINE)                  # 代码块结束围栏
    text = re.sub(r'(\*\*\*|\*\*|\*|__|_)(.+?)\1', r'\2', text)               # 粗体/斜体
    text = re.sub(r'`([^`]+)`', r'\1', text)                                  # 行内代码
    text = re.sub(r'^\s*[-*+]\s+', '- ', text, flags=re.MULTILINE)            # 无序列表符号统一
    text = re.sub(r'\n{3,}', '\n\n', text)                                    # 多余空行收敛
    return text.strip()


def _add_inline_runs(paragraph, text: str) -> None:
    """将一行文本中的 **粗体**、*斜体*、`行内代码` 解析成带格式的 run，写入 docx 段落。"""
    pos = 0
    for m in _INLINE_TOKEN_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        token = m.group(0)
        if token.startswith("***") and token.endswith("***"):
            r = paragraph.add_run(token[3:-3])
            r.bold = True
            r.italic = True
        elif token.startswith("**") and token.endswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
        elif token.startswith("`") and token.endswith("`"):
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
        else:  # 单个 * 或 _ 包裹 → 斜体
            r = paragraph.add_run(token[1:-1])
            r.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _markdown_to_docx_bytes_fallback(markdown_text: str) -> bytes:
    """轻量 Markdown → docx 转换（python-docx 手工解析，pandoc 不可用时的兜底方案）。
    支持：# 标题、-/*/+ 无序列表、数字. 有序列表、```代码块```、**粗体**、*斜体*、`行内代码`。
    其余按普通段落处理。目标是可读的 Word 文档，不追求 100% Markdown 语义还原，
    也不支持样式模板（如需标准/模板样式，请安装 pandoc 并配置 DOCX_REFERENCE_TEMPLATE）。"""
    doc = Document()
    lines = (markdown_text or "").splitlines()

    in_code_block = False
    code_lines = []

    def flush_code_block():
        if not code_lines:
            return
        p = doc.add_paragraph()
        r = p.add_run("\n".join(code_lines))
        r.font.name = "Consolas"
        r.font.size = Pt(9.5)
        p.paragraph_format.left_indent = Pt(18)
        code_lines.clear()

    for raw_line in lines:
        line = raw_line.rstrip("\n")
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                in_code_block = True
            continue
        if in_code_block:
            code_lines.append(line)
            continue

        if not stripped:
            doc.add_paragraph("")
            continue

        h_match = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if h_match:
            level = len(h_match.group(1))
            doc.add_heading(h_match.group(2).strip(), level=min(level, 9))
            continue

        ul_match = re.match(r'^[-*+]\s+(.*)$', stripped)
        if ul_match:
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(p, ul_match.group(1))
            continue

        ol_match = re.match(r'^\d+[.)]\s+(.*)$', stripped)
        if ol_match:
            p = doc.add_paragraph(style="List Number")
            _add_inline_runs(p, ol_match.group(1))
            continue

        p = doc.add_paragraph()
        _add_inline_runs(p, stripped)

    if in_code_block:
        flush_code_block()  # 未闭合的代码块围栏，尽量保留已收集内容

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _markdown_to_docx_pandoc(markdown_text: str) -> bytes:
    """用系统 pandoc 把 Markdown 转成 docx。
    比手工 python-docx 解析更"标准"：完整支持表格、脚注、嵌套列表、引用块等
    Markdown 语义，且可以通过 --reference-doc 套用外部模板的样式（标题/正文
    字体、字号、颜色、页边距等），从而生成统一风格的文档，而不是 pandoc 默认
    的朴素样式。失败或超时会抛异常，由调用方决定是否降级。"""
    if not _PANDOC_PATH:
        raise RuntimeError("pandoc not installed")

    cmd = [_PANDOC_PATH, "-f", "markdown", "-t", "docx"]
    if DOCX_REFERENCE_TEMPLATE and os.path.isfile(DOCX_REFERENCE_TEMPLATE):
        cmd.append(f"--reference-doc={DOCX_REFERENCE_TEMPLATE}")
    cmd += ["-o", "-"]  # 输出到 stdout

    proc = subprocess.run(
        cmd,
        input=(markdown_text or "").encode("utf-8"),
        capture_output=True,
        timeout=PANDOC_TIMEOUT,
    )
    if proc.returncode != 0:
        stderr = proc.stderr.decode("utf-8", errors="ignore").strip()
        raise RuntimeError(f"pandoc exited {proc.returncode}: {stderr[:500]}")
    if not proc.stdout:
        raise RuntimeError("pandoc produced empty output")
    return proc.stdout


def _markdown_to_docx_bytes(markdown_text: str) -> tuple:
    """Markdown → docx 统一入口。返回 (docx_bytes, engine_name)。
    优先走 pandoc（更标准、可套模板样式）；pandoc 未安装或转换失败时自动降级
    到内置的轻量 python-docx 解析器，保证 docx 输出始终可用。"""
    if _PANDOC_PATH:
        try:
            return _markdown_to_docx_pandoc(markdown_text), "pandoc"
        except Exception as e:
            app.logger.warning(f"pandoc conversion failed, falling back to python-docx: {e}")
    return _markdown_to_docx_bytes_fallback(markdown_text), "python-docx-fallback"


def _apply_output_format(result: dict, output_format: str, filename: str) -> dict:
    """根据 output_format 改写成功响应的 content，附加 output_filename/mime_type。
    docx 场景下会把 content 换成 file_base64（二进制），并移除 content 字段。"""
    base_name = os.path.splitext(os.path.basename(filename or "converted"))[0] or "converted"
    content = result.get("content", "")
    result["output_format"] = output_format

    if output_format == "md":
        result["output_filename"] = f"{base_name}.md"
        result["mime_type"] = "text/markdown"

    elif output_format == "txt":
        result["output_filename"] = f"{base_name}.txt"
        result["mime_type"] = "text/plain"
        result["content"] = _strip_markdown_light(content)

    elif output_format == "docx":
        try:
            docx_bytes, docx_engine = _markdown_to_docx_bytes(content)
        except Exception as e:
            result["success"] = False
            result["error"] = f"Failed to build docx: {e}"
            result.pop("content", None)
            return result
        result["output_filename"] = f"{base_name}.docx"
        result["mime_type"] = ("application/vnd.openxmlformats-officedocument"
                                ".wordprocessingml.document")
        result["file_base64"] = base64.b64encode(docx_bytes).decode("ascii")
        result["docx_engine"] = docx_engine  # "pandoc" 或 "python-docx-fallback"
        result.pop("content", None)  # docx 是二进制文件，不再返回纯文本 content

    return result


def _respond(result: dict, status: int, output_format: str, filename: str):
    """统一出口：成功响应按 output_format 转换后再 jsonify；失败响应原样透传。"""
    if status == 200 and isinstance(result, dict) and result.get("success") and "content" in result:
        result = _apply_output_format(result, output_format, filename)
        if not result.get("success", True):
            status = 500
    return jsonify(result), status


# ============================================================
#  路由：文档转换（统一入口 + 智能路由）
# ============================================================

@app.route("/convert", methods=["POST"])
def convert_document():
    """
    统一转换入口。
    根据文件扩展名自动路由：
    - 办公文档 → 本地 markitdown（crop 无效，会被忽略）
    - 图像文件 → Win11 PaddleOCR（crop 会先本地裁剪再转发）
    - PDF → pdfplumber 分类后路由（crop 会按文本型/扫描型分别处理）
    """
    if "file" not in request.files:
        return jsonify({"success": False, "error": "No file part in the request"}), 400

    file = request.files["file"]
    if file.filename == "":
        return jsonify({"success": False, "error": "No selected file"}), 400

    # ---- 解析可选的 crop 字段 ----
    crop = None
    crop_raw = request.form.get("crop")
    if crop_raw:
        try:
            crop = _parse_crop(crop_raw)
        except ValueError as e:
            return jsonify({"success": False, "error": f"Invalid crop: {e}"}), 400

    # ---- 解析可选的 output_format 字段（默认 md，兼容旧前端）----
    try:
        output_format = _validate_output_format(request.form.get("output_format"))
    except ValueError as e:
        return jsonify({"success": False, "error": str(e)}), 400

    # ---- 拒绝没有意义的同格式转换：源文件已是 .docx 时不允许再输出 docx ----
    # 走这条路径本质是 docx → markitdown 拆成 Markdown → 再用 pandoc/python-docx
    # 重新生成一份"标准样式"docx，原始排版、图片位置、表格样式必然丢失，
    # 对用户没有价值，直接在服务端拒绝（前端也会禁用该选项，这里是兜底）。
    _, _src_ext = os.path.splitext(file.filename)
    if _src_ext.lower() == ".docx" and output_format == "docx":
        return jsonify({
            "success": False,
            "error": ("源文件已是 .docx，选择 docx 作为输出格式没有意义：内容会先被 markitdown "
                      "拆解成 Markdown 再重新生成一份通用样式的 docx，原始排版、图片、表格样式"
                      "必然丢失。请直接使用原文件，或改选 md/txt 输出。"),
            "filename": file.filename,
        }), 400

    task_started()
    t_start = time.perf_counter()
    try:
        file_bytes = file.read()
        file_size_mb = len(file_bytes) / (1024 * 1024)

        if file_size_mb > MAX_FILE_SIZE_MB:
            return jsonify({
                "success": False,
                "error": f"File too large ({file_size_mb:.1f}MB). Max: {MAX_FILE_SIZE_MB}MB.",
            }), 413

        _, ext = os.path.splitext(file.filename)
        ext = ext.lower()
        if not ext:
            return jsonify({
                "success": False,
                "error": "File extension missing, cannot determine format",
            }), 400

        # ---- 路由决策 ----
        if ext in LOCAL_ONLY_EXTENSIONS:
            result, status = _convert_local(file_bytes, file.filename, ext, t_start)
            if crop is not None:
                result["crop_ignored"] = "Region crop is not supported for office formats"
            return _respond(result, status, output_format, file.filename)

        elif ext in OCR_ONLY_EXTENSIONS:
            if crop is not None:
                try:
                    file_bytes = _crop_image_bytes(file_bytes, crop)
                except Exception as e:
                    return jsonify({
                        "success": False,
                        "error": f"Failed to crop image: {e}",
                    }), 400
            result, status = _convert_ocr(file_bytes, file.filename, t_start)
            if crop is not None and status == 200:
                result["crop_applied"] = True
            return _respond(result, status, output_format, file.filename)

        elif ext == PDF_EXTENSION:
            pdf_type = _classify_pdf(file_bytes)

            if pdf_type == "text":
                if crop is not None:
                    content = _crop_pdf_text(file_bytes, crop)
                    elapsed_ms = round((time.perf_counter() - t_start) * 1000)
                    result = {
                        "success": True,
                        "filename": file.filename,
                        "content": content,
                        "engine": "pdfplumber_crop",
                        "routed_to": "local",
                        "crop_applied": True,
                        "elapsed_ms": elapsed_ms,
                    }
                    return _respond(result, 200, output_format, file.filename)
                result, status = _convert_local(file_bytes, file.filename, ext, t_start)
                return _respond(result, status, output_format, file.filename)

            else:  # scanned
                if crop is not None:
                    result, status = _convert_ocr_cropped_pdf(
                        file_bytes, file.filename, crop, t_start)
                    return _respond(result, status, output_format, file.filename)
                result, status = _convert_ocr(file_bytes, file.filename, t_start)
                return _respond(result, status, output_format, file.filename)

        else:
            return jsonify({
                "success": False,
                "error": f"Unsupported format '{ext}'. See GET /supported_formats for available formats.",
                "filename": file.filename,
            }), 400

    except Exception as e:
        return jsonify({
            "success": False,
            "error": f"Conversion failed: {str(e)}",
            "filename": file.filename,
        }), 500

    finally:
        task_finished()


def _convert_local(file_bytes: bytes, filename: str, ext: str,
                   t_start: float) -> tuple:
    """本地 markitdown 转换。返回 (dict, status)。"""
    file_stream = io.BytesIO(file_bytes)
    result = markitdown.convert_stream(file_stream, file_extension=ext)
    elapsed_ms = round((time.perf_counter() - t_start) * 1000)
    file_size_mb = len(file_bytes) / (1024 * 1024)

    return {
        "success": True,
        "filename": filename,
        "file_size_mb": round(file_size_mb, 2),
        "content": result.text_content,
        "engine": "markitdown",
        "routed_to": "local",
        "elapsed_ms": elapsed_ms,
    }, 200


def _convert_ocr(file_bytes: bytes, filename: str,
                 t_start: float) -> tuple:
    """转发到 Win11 OCR 节点。包含健康检查驱动的故障转移。返回 (dict, status)。"""
    win11_health = _check_win11_health()

    if win11_health is None:
        # Win11 不可达 — 尝试本地降级（仅 PDF；图片没有本地降级路径）
        _, ext = os.path.splitext(filename)
        ext = ext.lower()
        if ext == PDF_EXTENSION:
            try:
                return _convert_local(file_bytes, filename, ext, t_start)
            except Exception:
                pass
        return {
            "success": False,
            "error": "Win11 OCR node unreachable and no local fallback for this format",
            "filename": filename,
        }, 503

    result, status = _forward_to_win11(file_bytes, filename)
    elapsed_ms = round((time.perf_counter() - t_start) * 1000)

    if result.get("success"):
        result["routed_to"] = "win11_ocr"
        # Win11 返回的 lines 带 bbox，这里做几何启发式段落/标题重建；
        # 多页 PDF 场景（page 字段来自 Win11 自己逐页打的标）也在这一步统一处理。
        _apply_layout_reconstruction(result)
    result["route_elapsed_ms"] = elapsed_ms

    return result, status


# ============================================================
#  启动入口
# ============================================================

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port, debug=False)
